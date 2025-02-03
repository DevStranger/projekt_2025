from transformers import pipeline
from docx import Document
import os



# Inicjalizacja modelu do podsumowywania
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

def split_text(text, max_length=1000):

    words = text.split(" ")
    segments = []
    segment = []
    
    for word in words:
        segment.append(word)
        if len(" ".join(segment)) > max_length:
            segments.append(" ".join(segment))
            segment = []
    
    if segment:
        segments.append(" ".join(segment))
    
    return segments


def generate_summary(docx_path, summary_docx_path):
    """
    Tworzy podsumowanie treści pliku .docx i zapisuje je w nowym pliku.
    """
    try:
        if not os.path.exists(docx_path):
            print(f"Plik {docx_path} nie istnieje.")
            return

        doc = Document(docx_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])

        start_marker = "Transkrypcja wykładu:"
        end_marker = "Zrzuty ekranu:"

        if start_marker in full_text:
            full_text = full_text.split(start_marker, 1)[-1]  # Pobiera wszystko po "Transkrypcja wykładu:"

        if end_marker in full_text:
            full_text = full_text.split(end_marker, 1)[0]  # Usuwa wszystko po "Zrzuty ekranu:"

        if len(full_text.strip()) < 50:
            print(f"Plik {docx_path} zawiera za mało treści do podsumowania.")
            return

        # Dzielimy tekst na mniejsze fragmenty
        segments = split_text(full_text, max_length=1000)
        if not segments:
            print(f"Brak segmentów do podsumowania dla {docx_path}")
            return

        summarized_segments = []
        print(f"📄 Przetwarzanie pliku: {docx_path}")
        print(f"📜 Oryginalny tekst ({len(full_text)} znaków):\n{full_text[:1000]}...\n")  # Podgląd pierwszych 1000 znaków
        print(f"🔍 Podzielony na {len(segments)} segmentów")

        for i, segment in enumerate(segments):
            print(f"📝 Segment {i} ({len(segment)} znaków): {segment[:300]}...")  # Podgląd pierwszych 300 znaków segmentu
            if len(segment.strip()) == 0:
                print(f"Pominięto pusty segment {i} w {docx_path}")
                continue
            
            try:
                summary = summarizer(segment, max_length=500, min_length=200, do_sample=False)
                if summary and 'summary_text' in summary[0]:
                    summarized_segments.append(summary[0]['summary_text'])
                else:
                    print(f"Nie udało się wygenerować podsumowania dla segmentu {i} w {docx_path}")
            except Exception as e:
                print(f"Błąd w summarizerze dla segmentu {i}: {e}")

        if not summarized_segments:
            print(f"Nie wygenerowano żadnych podsumowań dla {docx_path}")
            return

        # Połączenie podsumowań w jeden dokument
        final_summary = "\n\n".join(summarized_segments)

        # Tworzenie nowego pliku z podsumowaniem
        summary_doc = Document()
        summary_doc.add_paragraph("Podsumowanie:", style='Heading 1')
        summary_doc.add_paragraph(final_summary)
        summary_doc.save(summary_docx_path)

        print(f"Zapisano podsumowanie do: {summary_docx_path}")

    except Exception as e:
        print(f"Błąd podczas generowania podsumowania dla {docx_path}: {e}")


