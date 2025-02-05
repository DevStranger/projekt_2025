import os
from docx import Document
from whisper import load_model
import torchaudio
from pydub import AudioSegment
from docx.shared import Inches
from pyannote.audio import Pipeline
from collections import defaultdict
from .summary import generate_summary

# Inicjalizacja modelu Whisper
token = "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"
model = load_model("base")


def save_transcription_to_docx(transcription, docx_path, screenshots_folder, speaker_stats=None):
    """
    Zapisuje transkrypcję i screenshoty w pliku .docx.
    """
    doc = Document()
    doc.add_paragraph("Transkrypcja wykładu:", style='Heading 1')
    doc.add_paragraph(transcription)
    doc.add_paragraph("Zrzuty ekranu:", style='Heading 1')

    if os.path.exists(screenshots_folder):
        for screenshot in sorted(os.listdir(screenshots_folder)):
            screenshot_path = os.path.join(screenshots_folder, screenshot)
            doc.add_picture(screenshot_path, width=Inches(5))
    else:
        doc.add_paragraph("Brak zrzutów ekranu.")

    if speaker_stats:
        doc.add_paragraph("Statystyki mówców:", style='Heading 1')
        for speaker, (duration, speed) in speaker_stats.items():
            doc.add_paragraph(f"{speaker}: {duration:.2f} sekund, {speed:.2f} słów/sek")

    doc.save(docx_path)
    print(f"Zapisano plik .docx: {docx_path}")


# Konwersja audio do 16kHz WAV

def load_and_convert_audio(audio_path):
    file_extension = os.path.splitext(audio_path)[1].lower()
    if file_extension == '.mp3':
        wav_path = audio_path.replace('.mp3', '.wav')
        audio = AudioSegment.from_mp3(audio_path)
        audio = audio.set_frame_rate(16000).set_channels(1)
        audio.export(wav_path, format="wav")
        audio_path = wav_path

    signal, sample_rate = torchaudio.load(audio_path)
    if sample_rate != 16000:
        signal = torchaudio.transforms.Resample(orig_freq=sample_rate, new_freq=16000)(signal)
    return signal, 16000


def diarize_audio(audio_path, token):
    """Diarizacja audio"""
    pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token=token)
    diarization = pipeline(audio_path)
    return diarization


def transcribe_audio(audio_path):
    signal, sample_rate = load_and_convert_audio(audio_path)
    transcription = model.transcribe(audio_path)
    return transcription['segments']


def combine_transcription_with_diarization(transcription_segments, diarization):
    diarized_text = []
    speaker_durations = defaultdict(float)
    speaker_word_counts = defaultdict(int)
    current_speaker = None
    speaker_transcriptions = defaultdict(list)

    for segment in transcription_segments:
        start_time, end_time, text = segment['start'], segment['end'], segment['text']
        word_count = len(text.split())

        speaker = None
        for turn, _, speaker_id in diarization.itertracks(yield_label=True):
            if turn.start <= start_time <= turn.end or turn.start <= end_time <= turn.end:
                speaker = speaker_id
                break

        if speaker is None:
            continue

        speaker_durations[speaker] += (end_time - start_time)
        speaker_word_counts[speaker] += word_count
        speaker_transcriptions[speaker].append(text)

    for speaker, sentences in speaker_transcriptions.items():
        diarized_text.append(f"{speaker}:\n{' '.join(sentences)}")

    speaker_speeds = {spk: speaker_word_counts[spk] / speaker_durations[spk] for spk in speaker_durations}
    return "\n".join(diarized_text), speaker_durations, speaker_speeds


def process_audio_and_save_transcription(audio_folder, docx_folder, screenshots_folder):
    if not os.path.exists(docx_folder):
        os.makedirs(docx_folder)

    for filename in os.listdir(audio_folder):
        if filename.endswith(".mp3") or filename.endswith(".wav"):
            audio_file_path = os.path.join(audio_folder, filename)
            docx_filename = os.path.join(docx_folder, f"{os.path.splitext(filename)[0]}.docx")
            raw_docx_filename = os.path.join(docx_folder, f"raw_{os.path.splitext(filename)[0]}.docx")
            summary_filename = os.path.join(docx_folder, f"summary_{os.path.splitext(filename)[0]}.docx")

            if os.path.exists(docx_filename):
                print(f"Pominięto {filename}, transkrypcja już istnieje.")
                continue

            try:
                transcription_segments = transcribe_audio(audio_file_path)
                raw_transcription_text = " ".join([seg['text'] for seg in transcription_segments])
                save_transcription_to_docx(raw_transcription_text, raw_docx_filename, screenshots_folder)

                diarization = diarize_audio(audio_file_path, token)
                diarized_text, speaker_durations, speaker_speeds = combine_transcription_with_diarization(
                    transcription_segments, diarization)

                speaker_stats = {spk: (speaker_durations[spk], speaker_speeds[spk]) for spk in speaker_durations}

                save_transcription_to_docx(diarized_text, docx_filename, screenshots_folder, speaker_stats)
                generate_summary(raw_docx_filename, summary_filename)

                print(f"Zapisano transkrypcję i podsumowanie do: {docx_filename}, {summary_filename}")
            except Exception as e:
                print(f"Błąd przy przetwarzaniu {filename}: {e}")


if __name__ == "__main__":
    audio_folder = r"recordings"
    docx_folder = os.path.join(os.getcwd(), 'recordings', 'notes')
    screenshots_folder = os.path.join(os.getcwd(), 'recordings', 'screenshots')

    if not os.path.exists(docx_folder):
        os.makedirs(docx_folder)

    process_audio_and_save_transcription(audio_folder, docx_folder, screenshots_folder)
